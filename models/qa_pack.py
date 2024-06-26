import pdfkit

from models.user import User
from models.tag import Tag
from utils import ParagraphExt

from firebase_config import firebase_db
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import render_template


class QA():
    def __init__(self, 
                 question: str, 
                 answer: str, 
                 id:str = None):
        self.id = id
        self.question = question
        self.answer = answer


class QAPack():
    def __init__(self, 
                 name: str, 
                 qas: list[QA] = [], 
                 tags_id: list[str] = [], 
                 total_qas: int = 0,
                 id: str = None, 
                 owner_id:str = None):
        
        self.id = id
        self.owner_id = owner_id
        self.name = name
        self.qas = qas
        self.tags_id = tags_id
        self.total_qas = total_qas


    def save(self, user: User):
        _, doc_ref = firebase_db.collection('qa_packs').add({
            'name': self.name,
            'owner_id': user.id,
            'tags_id': self.tags_id,
            'total_qas': self.total_qas
        })

        self.id = doc_ref.id
        self.owner_id = user.id

        return self


    def save_qas(self):
        for qa in self.qas:
            _, doc_ref = firebase_db.collection('qa_packs').document(self.id).collection('qas').add({
                'question': qa.question,
                'answer': qa.answer
            })

            qa.id = doc_ref.id
            self.total_qas += 1
        
        firebase_db.collection('qa_packs').document(self.id).update({
            'total_qas': self.total_qas
        })

        return self


    def update(self):
        firebase_db.collection('qa_packs').document(self.id).update({
            'name': self.name,
            'tags_id': self.tags_id
        })

        return self


    def update_qas(self):
        current_qas = firebase_db.collection('qa_packs').document(self.id).collection('qas').get()
        current_qa_ids = [qa.id for qa in current_qas]

        self.total_qas = len(current_qa_ids)

        for qa in self.qas:
            if qa.id is None or qa.id not in current_qa_ids or qa.id == '':
                _, doc_ref = firebase_db.collection('qa_packs').document(self.id).collection('qas').add({
                    'question': qa.question,
                    'answer': qa.answer
                })

                qa.id = doc_ref.id
                self.total_qas += 1
            
            else:
                firebase_db.collection('qa_packs').document(self.id).collection('qas').document(qa.id).update({
                    'question': qa.question,
                    'answer': qa.answer
                })
        
        for qa_id in current_qa_ids:
            if qa_id not in [qa.id for qa in self.qas]:
                firebase_db.collection('qa_packs').document(self.id).collection('qas').document(qa_id).delete()
                self.total_qas -= 1
        
        firebase_db.collection('qa_packs').document(self.id).update({
            'total_qas': self.total_qas
        })

        return self
    

    def delete(self):
        qas = firebase_db.collection('qa_packs').document(self.id).collection('qas')

        for qa in qas.stream():
            qa.reference.delete()

        firebase_db.collection('qa_packs').document(self.id).delete()

        return self
    

    def export_docx(self):
        doc = Document()
        
        title = doc.add_heading(self.name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('')

        for qa in self.qas:
            doc.add_paragraph(qa.question, style='List Number')

        doc.add_page_break()

        answer_title = doc.add_heading("Kunci Jawaban", level=1)
        answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('')

        for index, qa in enumerate(self.qas):
            p = doc.add_paragraph(qa.answer, style='List Number')
            if index == 0:
                ParagraphExt(p).restart_numbering()

        doc.save('./export/export.docx')

        return './export/export.docx'
    

    def export_pdf(self):
        html = render_template('export_template/pdf.html', qa_pack=self)
        pdf_path = './export/export.pdf'
        pdfkit.from_string(html, pdf_path)

        return pdf_path
    

    def export_gift(self):
        gift = ''

        for qa in self.qas:
            gift += qa.question + ' {=' + qa.answer + '}\n\n'

        with open('./export/export.txt', 'w') as file:
            file.write(gift)

        return './export/export.txt'


    @staticmethod
    def get_all(user: User):
        qa_packs = firebase_db.collection('qa_packs').where('owner_id', '==', user.id).stream()
        return [QAPack(qa_pack.to_dict()['name'], id=qa_pack.id, owner_id=qa_pack.to_dict()['owner_id'], tags_id=qa_pack.to_dict()['tags_id'], total_qas=qa_pack.to_dict()['total_qas']) for qa_pack in qa_packs]
    

    @staticmethod
    def get_all_for_table(user: User):
        qa_packs = firebase_db.collection('qa_packs').where('owner_id', '==', user.id).stream()

        table_data = []
        tags = {}

        for qa_pack in qa_packs:
            qa_pack_id = qa_pack.id
            qa_pack_dict = qa_pack.to_dict()

            qa_pack_tags = []
            for tag_id in qa_pack_dict['tags_id']:
                if tag_id not in tags:
                    tag = firebase_db.collection('tags').document(tag_id).get().to_dict()
                    tags[tag_id] = { 'name': tag['name'], 'color': tag['color']}
                qa_pack_tags.append(tags[tag_id])

            table_data.append({
                'id': qa_pack_id,
                'name': qa_pack_dict['name'],
                'tags': qa_pack_tags,
                'total_qas': qa_pack_dict['total_qas']
            })

        return table_data
    

    @staticmethod
    def get_all_for_table_admin():
        qa_packs = firebase_db.collection('qa_packs').stream()
        users = firebase_db.collection('users').stream()

        user_dict = {}
        for user in users:
            user_dict[user.id] = user.to_dict()['name']

        table_data = []
        tags = {}

        for qa_pack in qa_packs:
            qa_pack_id = qa_pack.id
            qa_pack_dict = qa_pack.to_dict()

            table_data.append({
                'id': qa_pack_id,
                'name': qa_pack_dict['name'],
                'total_qas': qa_pack_dict['total_qas'],
                'owner_name': user_dict[qa_pack_dict['owner_id']],
            })

        return table_data
    

    @staticmethod
    def get_by_id(qa_pack_id: str):
        qa_pack = firebase_db.collection('qa_packs').document(qa_pack_id).get().to_dict()
        qas = firebase_db.collection('qa_packs').document(qa_pack_id).collection('qas').stream()
        tags = [Tag.get_by_id(tag_id) for tag_id in qa_pack['tags_id']]

        return QAPack(qa_pack['name'], id=qa_pack_id, owner_id=qa_pack['owner_id'], tags_id=qa_pack['tags_id'], total_qas=qa_pack['total_qas'], qas=[QA(qa.to_dict()['question'], qa.to_dict()['answer'], id=qa.id) for qa in qas]), tags
    

    @staticmethod
    def get_by_id_no_tags(qa_pack_id: str):
        qa_pack = firebase_db.collection('qa_packs').document(qa_pack_id).get().to_dict()
        qas = firebase_db.collection('qa_packs').document(qa_pack_id).collection('qas').stream()

        return QAPack(qa_pack['name'], id=qa_pack_id, owner_id=qa_pack['owner_id'], tags_id=qa_pack['tags_id'], total_qas=qa_pack['total_qas'], qas=[QA(qa.to_dict()['question'], qa.to_dict()['answer'], id=qa.id) for qa in qas])